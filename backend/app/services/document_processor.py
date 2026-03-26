"""Text extraction and chunking (no LLM / vector store)."""

from __future__ import annotations

import logging
from typing import List

import PyPDF2
from docx import Document as DocxDocument
from langchain.schema import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter

logger = logging.getLogger(__name__)

_DEFAULT_CHUNK_SIZE = 1000
_DEFAULT_CHUNK_OVERLAP = 200


class DocumentProcessor:
    """Load local files and split into LangChain Documents (metadata: source only)."""

    @staticmethod
    def process_file(file_path: str, file_type: str) -> List[Document]:
        try:
            ext = file_type.lower()
            if ext == "pdf":
                text = DocumentProcessor._extract_pdf_text(file_path)
            elif ext in ("docx", "doc"):
                text = DocumentProcessor._extract_docx_text(file_path)
            else:
                text = DocumentProcessor._extract_text_file(file_path)

            splitter = RecursiveCharacterTextSplitter(
                chunk_size=_DEFAULT_CHUNK_SIZE,
                chunk_overlap=_DEFAULT_CHUNK_OVERLAP,
                length_function=len,
            )
            chunks = splitter.split_text(text)
            return [
                Document(page_content=chunk, metadata={"source": file_path})
                for chunk in chunks
            ]
        except Exception as e:
            logger.error("Error processing document: %s", e)
            raise

    @staticmethod
    def _extract_pdf_text(file_path: str) -> str:
        text = ""
        with open(file_path, "rb") as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
        return text

    @staticmethod
    def _extract_docx_text(file_path: str) -> str:
        doc = DocxDocument(file_path)
        parts = []
        for paragraph in doc.paragraphs:
            parts.append(paragraph.text)
        return "\n".join(parts)

    @staticmethod
    def _extract_text_file(file_path: str) -> str:
        with open(file_path, "r", encoding="utf-8", errors="replace") as file:
            return file.read()
