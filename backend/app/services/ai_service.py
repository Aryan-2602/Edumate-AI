import os
import json
from typing import List, Dict, Any, Optional
from langchain_openai import OpenAIEmbeddings, ChatOpenAI
from langchain_community.vectorstores import Chroma
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain.schema import Document
from langchain.chains import RetrievalQA
from langchain.prompts import PromptTemplate
import PyPDF2
from docx import Document as DocxDocument
from app.config import settings
import wandb
import logging

logger = logging.getLogger(__name__)


class AIService:
    def __init__(self):
        self.embeddings = OpenAIEmbeddings(
            openai_api_key=settings.openai_api_key,
            model="text-embedding-ada-002"
        )
        self.llm = ChatOpenAI(
            openai_api_key=settings.openai_api_key,
            model=settings.openai_model,
            max_tokens=settings.openai_max_tokens,
            temperature=settings.openai_temperature
        )
        self.vector_store = None
        self._init_wandb()
    
    def _init_wandb(self):
        """Initialize Weights & Biases for tracking"""
        if settings.wandb_api_key:
            try:
                wandb.login(key=settings.wandb_api_key)
                wandb.init(project=settings.wandb_project)
            except Exception as e:
                logger.warning(f"Failed to initialize W&B: {e}")
    
    def process_document(self, file_path: str, file_type: str) -> List[Document]:
        """Process uploaded document and split into chunks"""
        try:
            # Extract text based on file type
            if file_type.lower() == 'pdf':
                text = self._extract_pdf_text(file_path)
            elif file_type.lower() in ['docx', 'doc']:
                text = self._extract_docx_text(file_path)
            else:
                text = self._extract_text_file(file_path)
            
            # Split text into chunks
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=1000,
                chunk_overlap=200,
                length_function=len,
            )
            chunks = text_splitter.split_text(text)
            
            # Convert to LangChain Document objects
            documents = [Document(page_content=chunk, metadata={"source": file_path}) for chunk in chunks]
            
            # Track processing in W&B
            if wandb.run:
                wandb.log({
                    "document_processed": True,
                    "file_type": file_type,
                    "chunk_count": len(chunks),
                    "total_text_length": len(text)
                })
            
            return documents
            
        except Exception as e:
            logger.error(f"Error processing document: {e}")
            raise
    
    def _extract_pdf_text(self, file_path: str) -> str:
        """Extract text from PDF file"""
        text = ""
        with open(file_path, 'rb') as file:
            pdf_reader = PyPDF2.PdfReader(file)
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
        return text
    
    def _extract_docx_text(self, file_path: str) -> str:
        """Extract text from DOCX file"""
        doc = DocxDocument(file_path)
        text = ""
        for paragraph in doc.paragraphs:
            text += paragraph.text + "\n"
        return text
    
    def _extract_text_file(self, file_path: str) -> str:
        """Extract text from plain text file"""
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    def create_embeddings(self, documents: List[Document], collection_name: str) -> Chroma:
        """Create embeddings and store in Chroma vector database"""
        try:
            # Create vector store
            vector_store = Chroma.from_documents(
                documents=documents,
                embedding=self.embeddings,
                persist_directory=settings.chroma_persist_directory,
                collection_name=collection_name
            )
            
            # Persist to disk
            vector_store.persist()
            
            # Track embeddings creation
            if wandb.run:
                wandb.log({
                    "embeddings_created": True,
                    "collection_name": collection_name,
                    "document_count": len(documents)
                })
            
            return vector_store
            
        except Exception as e:
            logger.error(f"Error creating embeddings: {e}")
            raise
    
    def answer_question(self, question: str, collection_name: str, top_k: int = 5) -> Dict[str, Any]:
        """Answer a question using RAG with the specified collection"""
        try:
            # Load vector store
            vector_store = Chroma(
                persist_directory=settings.chroma_persist_directory,
                embedding_function=self.embeddings,
                collection_name=collection_name
            )
            
            # Create retrieval QA chain
            qa_chain = RetrievalQA.from_chain_type(
                llm=self.llm,
                chain_type="stuff",
                retriever=vector_store.as_retriever(search_kwargs={"k": top_k}),
                return_source_documents=True
            )
            
            # Generate answer
            result = qa_chain({"query": question})
            
            # Extract source documents
            source_docs = []
            if result.get("source_documents"):
                for doc in result["source_documents"]:
                    source_docs.append({
                        "content": doc.page_content[:200] + "...",
                        "metadata": doc.metadata
                    })
            
            response = {
                "answer": result["result"],
                "sources": source_docs,
                "question": question
            }
            
            # Track Q&A in W&B
            if wandb.run:
                wandb.log({
                    "question_answered": True,
                    "question_length": len(question),
                    "answer_length": len(result["result"]),
                    "source_count": len(source_docs)
                })
            
            return response
            
        except Exception as e:
            logger.error(f"Error answering question: {e}")
            raise
    
    def generate_quiz(self, content: str, num_questions: int = 5) -> List[Dict[str, Any]]:
        """Generate quiz questions from content"""
        try:
            prompt_template = PromptTemplate(
                input_variables=["content", "num_questions"],
                template="""
                Based on the following content, generate {num_questions} multiple choice questions.
                Each question should have 4 options (A, B, C, D) with only one correct answer.
                
                Content: {content}
                
                Format your response as a JSON array with the following structure:
                [
                    {{
                        "question": "Question text here?",
                        "options": ["Option A", "Option B", "Option C", "Option D"],
                        "correct_answer": "A",
                        "explanation": "Explanation of why this is correct"
                    }}
                ]
                
                Make sure the questions are educational and test understanding of the content.
                """
            )
            
            prompt = prompt_template.format(content=content, num_questions=num_questions)
            response = self.llm.invoke(prompt)
            
            # Parse JSON response
            try:
                quiz_questions = json.loads(response.content)
            except json.JSONDecodeError:
                # Fallback: try to extract JSON from the response
                content = response.content
                start = content.find('[')
                end = content.rfind(']') + 1
                if start != -1 and end != 0:
                    quiz_questions = json.loads(content[start:end])
                else:
                    raise ValueError("Could not parse quiz questions from LLM response")
            
            # Track quiz generation
            if wandb.run:
                wandb.log({
                    "quiz_generated": True,
                    "question_count": len(quiz_questions),
                    "source_content_length": len(content)
                })
            
            return quiz_questions
            
        except Exception as e:
            logger.error(f"Error generating quiz: {e}")
            raise
    
    def generate_flashcards(self, content: str, num_cards: int = 10) -> List[Dict[str, str]]:
        """Generate flashcards from content"""
        try:
            prompt_template = PromptTemplate(
                input_variables=["content", "num_cards"],
                template="""
                Based on the following content, generate {num_cards} flashcards.
                Each flashcard should have a front (question/concept) and back (answer/explanation).
                
                Content: {content}
                
                Format your response as a JSON array with the following structure:
                [
                    {{
                        "front": "Front of the card (question or concept)",
                        "back": "Back of the card (answer or explanation)"
                    }}
                ]
                
                Make the flashcards concise but informative, covering key concepts from the content.
                """
            )
            
            prompt = prompt_template.format(content=content, num_cards=num_cards)
            response = self.llm.invoke(prompt)
            
            # Parse JSON response
            try:
                flashcards = json.loads(response.content)
            except json.JSONDecodeError:
                # Fallback: try to extract JSON from the response
                content = response.content
                start = content.find('[')
                end = content.rfind(']') + 1
                if start != -1 and end != 0:
                    flashcards = json.loads(content[start:end])
                else:
                    raise ValueError("Could not parse flashcards from LLM response")
            
            # Track flashcard generation
            if wandb.run:
                wandb.log({
                    "flashcards_generated": True,
                    "card_count": len(flashcards),
                    "source_content_length": len(content)
                })
            
            return flashcards
            
        except Exception as e:
            logger.error(f"Error generating flashcards: {e}")
            raise
    
    def get_similar_documents(self, query: str, collection_name: str, top_k: int = 5) -> List[Dict[str, Any]]:
        """Find similar documents using semantic search"""
        try:
            # Load vector store
            vector_store = Chroma(
                persist_directory=settings.chroma_persist_directory,
                embedding_function=self.embeddings,
                collection_name=collection_name
            )
            
            # Perform similarity search
            results = vector_store.similarity_search_with_score(query, k=top_k)
            
            similar_docs = []
            for doc, score in results:
                similar_docs.append({
                    "content": doc.page_content,
                    "metadata": doc.metadata,
                    "similarity_score": float(score)
                })
            
            return similar_docs
            
        except Exception as e:
            logger.error(f"Error finding similar documents: {e}")
            raise
